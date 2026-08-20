"""读一份 .docx，把所有段落带序号打印出来。用途：验货，看看 python-docx 眼里的文档长什么样。"""

import sys
from docx import Document

# 从命令行拿文件路径：python read_docx.py 合同.docx
path = sys.argv[1]

# Document() 把整个 .docx 解析成一个对象
doc = Document(path)

# enumerate 给每个段落配一个序号 i，从 0 开始
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()   # strip() 去掉首尾空白
    if not text:               # 空段落跳过不打印，但序号 i 不会因此错位
        continue
    print(i, "|", text)

print("---")
print("段落总数（含空段落）:", len(doc.paragraphs))
print("表格数量:", len(doc.tables))
